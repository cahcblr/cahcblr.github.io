# /// script
# dependencies = [
#   "python-docx",
#   "pyyaml",
# ]
# ///

import os
import re
import yaml
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

# --- Configurations ---
OUTPUT_DIR = "tmp"
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "sanskrit_review.docx")
DATA_DIR = "_data"
POSTS_DIR = "_posts"

# --- Style Utilities ---
def set_cell_background(cell, color_hex):
    """Sets the background color of a table cell."""
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets the internal margins (padding) of a table cell (values in dxas)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def format_paragraph(paragraph, font_name="Georgia", font_size=11, bold=False, italic=False, color_rgb=(51, 51, 51)):
    """Applies font formatting to a paragraph."""
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.15
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.bold = bold
        run.font.italic = italic
        run.font.color.rgb = RGBColor(*color_rgb)

def add_styled_heading(doc, text, level):
    """Adds a styled heading to the document."""
    heading = doc.add_heading(text, level=level)
    heading.paragraph_format.space_before = Pt(16)
    heading.paragraph_format.space_after = Pt(6)
    color = (26, 54, 93) if level == 1 else (44, 82, 130)
    for run in heading.runs:
        run.font.name = "Georgia"
        run.font.size = Pt(18 if level == 1 else 14)
        run.bold = True
        run.font.color.rgb = RGBColor(*color)
    return heading

def create_styled_table(doc, headers, col_widths):
    """Creates a table with styled headers."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    
    # Format headers
    hdr_cells = table.rows[0].cells
    for i, header_text in enumerate(headers):
        hdr_cells[i].text = header_text
        hdr_cells[i].width = col_widths[i]
        set_cell_background(hdr_cells[i], "1A365D")  # Dark Blue
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=150, right=150)
        p = hdr_cells[i].paragraphs[0]
        p.runs[0].bold = True
        format_paragraph(p, font_name="Georgia", font_size=10, bold=True, color_rgb=(255, 255, 255))
        
    return table

def add_table_row(table, data, col_widths, is_alternate=False):
    """Adds a row to the table with optional alternating background shading."""
    row_cells = table.add_row().cells
    bg_color = "F7FAFC" if is_alternate else "FFFFFF"
    
    for i, text in enumerate(data):
        row_cells[i].text = str(text or "")
        row_cells[i].width = col_widths[i]
        set_cell_background(row_cells[i], bg_color)
        set_cell_margins(row_cells[i], top=100, bottom=100, left=150, right=150)
        format_paragraph(row_cells[i].paragraphs[0], font_name="Georgia", font_size=10)

# --- Parser Functions ---
def clean_markdown(text):
    """Remove markdown bold/italic/link syntax for clean reading."""
    if not text:
        return ""
    text = re.sub(r'\*\*([^*]+)\*\*|__([^_]+)__', r'\1\2', text)
    text = re.sub(r'\*([^*]+)\*|_([^_]+)_', r'\1\2', text)
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
    return text.strip()

def extract_bilingual_post_blocks(filepath):
    """Extracts English/Sanskrit block pairs from posts."""
    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Matches lang-blocks containing en and sa structures
    pattern = r'<div\s+class="lang-content"\s+lang="en"[^>]*>(.*?)<\/div>\s*<div\s+class="lang-content"\s+lang="sa"[^>]*>(.*?)<\/div>'
    matches = re.findall(pattern, content, re.DOTALL)
    
    blocks = []
    for en_html, sa_html in matches:
        # Strip style and script tags along with their inner contents
        en_html = re.sub(r'<style\b[^>]*>([\s\S]*?)<\/style>', '', en_html, flags=re.I)
        en_html = re.sub(r'<script\b[^>]*>([\s\S]*?)<\/script>', '', en_html, flags=re.I)
        sa_html = re.sub(r'<style\b[^>]*>([\s\S]*?)<\/style>', '', sa_html, flags=re.I)
        sa_html = re.sub(r'<script\b[^>]*>([\s\S]*?)<\/script>', '', sa_html, flags=re.I)

        # Strip internal tags
        en_text = re.sub(r'<[^>]+>', '', en_html)
        sa_text = re.sub(r'<[^>]+>', '', sa_html)
        blocks.append((clean_markdown(en_text), clean_markdown(sa_text)))
        
    return blocks

# --- Main Document Builder ---
def main():
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    doc = Document()
    
    # Set document margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Document Header
    title_p = doc.add_paragraph()
    title_run = title_p.add_run("CAHC Sanskrit Translation Review Catalog")
    title_run.font.name = "Georgia"
    title_run.font.size = Pt(24)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor(26, 54, 93)
    title_p.paragraph_format.space_after = Pt(4)
    
    subtitle_p = doc.add_paragraph()
    subtitle_run = subtitle_p.add_run("Use Suggesting Mode (Track Changes) in Google Docs to propose translation updates.")
    subtitle_run.font.name = "Georgia"
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.italic = True
    subtitle_run.font.color.rgb = RGBColor(113, 128, 150)
    subtitle_p.paragraph_format.space_after = Pt(24)

    # --- SECTION 1: Page UI Strings & Short Labels ---
    add_styled_heading(doc, "Section 1: General Website UI Labels", level=1)
    
    ui_headers = ["Page / Section", "English Label", "Sanskrit (Current)", "Reviewer Suggestion / Note"]
    ui_col_widths = [Inches(1.5), Inches(2.0), Inches(2.0), Inches(2.0)]
    ui_table = create_styled_table(doc, ui_headers, ui_col_widths)
    
    # 1a. Load home page translations
    home_file = os.path.join(DATA_DIR, "pages", "home.yml")
    if os.path.exists(home_file):
        with open(home_file, 'r', encoding='utf-8') as f:
            home_data = yaml.safe_load(f)
        
        # Add book status row
        bullets = home_data.get("pub_bullets", [])
        if len(bullets) > 1:
            add_table_row(ui_table, [
                "Home / Publication Facts", 
                clean_markdown(bullets[1].get("en")), 
                clean_markdown(bullets[1].get("sa")), 
                ""
            ], ui_col_widths)

        # Add caption rows
        add_table_row(ui_table, [
            "Home / IKS Caption", 
            home_data.get("iks_caption", {}).get("en", ""), 
            home_data.get("iks_caption", {}).get("sa", ""), 
            ""
        ], ui_col_widths, is_alternate=True)

    # 1b. Load people profiles links & metadata
    people_file = os.path.join(DATA_DIR, "people.yml")
    if os.path.exists(people_file):
        with open(people_file, 'r', encoding='utf-8') as f:
            people_data = yaml.safe_load(f)
        
        alt = False
        for idx, person in enumerate(people_data):
            name_en = person.get("name", {}).get("en", "")
            role_en = person.get("role", {}).get("en", "")
            role_sa = person.get("role", {}).get("sa", "")
            add_table_row(ui_table, [
                f"People / Role ({name_en})",
                role_en,
                role_sa,
                ""
            ], ui_col_widths, is_alternate=alt)
            alt = not alt

    doc.add_page_break()

    # --- SECTION 2: Long-Form Prose & Biographies (Side-by-Side) ---
    add_styled_heading(doc, "Section 2: Long-Form Bios & Static Page Paragraphs", level=1)
    
    prose_headers = ["English Content", "Sanskrit (Current Version - Edit Directly here)"]
    prose_col_widths = [Inches(3.75), Inches(3.75)]

    # 2a. People Biographies
    if os.path.exists(people_file):
        add_styled_heading(doc, "Biographies & Leadership Profiles", level=2)
        for person in people_data:
            name_en = person.get("name", {}).get("en", "")
            bio_en = person.get("bio", {}).get("en", "")
            bio_sa = person.get("bio", {}).get("sa", "")
            
            if bio_en or bio_sa:
                p_label = doc.add_paragraph()
                p_label.add_run(f"Profile: {name_en}").bold = True
                p_label.paragraph_format.space_before = Pt(8)
                p_label.paragraph_format.space_after = Pt(2)
                
                table = create_styled_table(doc, prose_headers, prose_col_widths)
                
                # Split bios by paragraphs
                en_paras = [clean_markdown(p) for p in (bio_en or "").split("\n\n") if p.strip()]
                sa_paras = [clean_markdown(p) for p in (bio_sa or "").split("\n\n") if p.strip()]
                
                max_paras = max(len(en_paras), len(sa_paras))
                for idx in range(max_paras):
                    en_text = en_paras[idx] if idx < len(en_paras) else ""
                    sa_text = sa_paras[idx] if idx < len(sa_paras) else ""
                    add_table_row(table, [en_text, sa_text], prose_col_widths, is_alternate=(idx % 2 == 1))

    # 2b. About Page Content
    about_file = os.path.join(DATA_DIR, "pages", "about.yml")
    if os.path.exists(about_file):
        add_styled_heading(doc, "About Page Static Content", level=2)
        with open(about_file, 'r', encoding='utf-8') as f:
            about_data = yaml.safe_load(f)
            
        about_table = create_styled_table(doc, prose_headers, prose_col_widths)
        add_table_row(about_table, [
            clean_markdown(about_data.get("intro_p", {}).get("en", "")),
            clean_markdown(about_data.get("intro_p", {}).get("sa", ""))
        ], prose_col_widths)
        
        # Add leadership section intro
        add_table_row(about_table, [
            clean_markdown(about_data.get("leadership", {}).get("p1", "").get("en", "")),
            clean_markdown(about_data.get("leadership", {}).get("p1", "").get("sa", ""))
        ], prose_col_widths, is_alternate=True)

    doc.add_page_break()

    # --- SECTION 3: Books & Publications Catalog ---
    add_styled_heading(doc, "Section 3: Books & Publications Catalog", level=1)
    
    books_file = os.path.join(DATA_DIR, "books.yml")
    if os.path.exists(books_file):
        with open(books_file, 'r', encoding='utf-8') as f:
            books_data = yaml.safe_load(f)
            
        for book in books_data:
            title_en = book.get("title", {}).get("en", "")
            desc_en = book.get("description", {}).get("en", "")
            desc_sa = book.get("description", {}).get("sa", "")
            
            p_label = doc.add_paragraph()
            p_label.add_run(f"Book: {title_en}").bold = True
            p_label.paragraph_format.space_before = Pt(8)
            p_label.paragraph_format.space_after = Pt(2)
            
            book_table = create_styled_table(doc, prose_headers, prose_col_widths)
            
            # Row 1: Title & Author
            add_table_row(book_table, [
                f"Title: {title_en}\nAuthor: {book.get('author', {}).get('en', '')}",
                f"शीर्षकम्: {book.get('title', {}).get('sa', '')}\nलेखकः: {book.get('author', {}).get('sa', '')}"
            ], prose_col_widths)
            
            # Row 2: Subtitle
            add_table_row(book_table, [
                f"Subtitle: {book.get('subtitle', {}).get('en', '')}",
                f"उपशीर्षकम्: {book.get('subtitle', {}).get('sa', '')}"
            ], prose_col_widths, is_alternate=True)
            
            # Row 3: Description
            add_table_row(book_table, [
                clean_markdown(desc_en),
                clean_markdown(desc_sa)
            ], prose_col_widths)

    doc.add_page_break()

    # --- SECTION 4: News & Announcements (Blog Posts) ---
    add_styled_heading(doc, "Section 4: News & Outreach Announcements", level=1)
    
    if os.path.exists(POSTS_DIR):
        posts = sorted([f for f in os.listdir(POSTS_DIR) if f.endswith(".markdown")], reverse=True)
        for post_filename in posts:
            post_path = os.path.join(POSTS_DIR, post_filename)
            blocks = extract_bilingual_post_blocks(post_path)
            
            if blocks:
                # Extract date/title from filename
                date_match = re.match(r'(\d{4}-\d{2}-\d{2})-(.*)\.markdown', post_filename)
                date_str = date_match.group(1) if date_match else "Date unknown"
                title_slug = date_match.group(2).replace("-", " ").title() if date_match else post_filename
                
                p_label = doc.add_paragraph()
                p_label.add_run(f"Announcement ({date_str}): {title_slug}").bold = True
                p_label.paragraph_format.space_before = Pt(8)
                p_label.paragraph_format.space_after = Pt(2)
                
                post_table = create_styled_table(doc, prose_headers, prose_col_widths)
                for idx, (en_text, sa_text) in enumerate(blocks):
                    add_table_row(post_table, [en_text, sa_text], prose_col_widths, is_alternate=(idx % 2 == 1))

    # Save the generated document
    doc.save(OUTPUT_FILE)
    print(f"Success: Translation review document generated at '{OUTPUT_FILE}'!")

if __name__ == "__main__":
    main()
